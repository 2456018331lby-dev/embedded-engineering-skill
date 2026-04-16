#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
doc_output_mcp.py

MCP server for generating professional engineering documents from
embedded-engineering script outputs.

Tools provided
  doc_rf_design_report    Generate RF design report from calc_* JSON outputs
  doc_power_tree_report   Generate power architecture report from gen_power_tree output
  doc_project_summary     Generate full project summary document
  doc_export_markdown     Convert any markdown text to a formatted Word document

All tools write .docx files to the configured output directory and return
the absolute file path so the user can open or share the document directly.

Transport: stdio (for Claude Code / local API environment)

Setup
  pip install mcp python-docx pydantic
  Add to .mcp.json — see project root config file.
"""

from __future__ import annotations

import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm
from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, ConfigDict, Field

# ---------------------------------------------------------------------------
# Server initialisation
# ---------------------------------------------------------------------------

mcp = FastMCP("doc_output_mcp")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_OUTPUT_DIR = str(Path.home() / "embedded-engineering" / "docs")

# Brand colours (engineering-neutral palette)
COLOR_TITLE     = RGBColor(0x1F, 0x49, 0x7D)   # deep blue
COLOR_HEADING   = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
COLOR_ACCENT    = RGBColor(0x70, 0xAD, 0x47)   # green (PASS)
COLOR_WARN      = RGBColor(0xFF, 0xC0, 0x00)   # amber (WARN)
COLOR_FAIL      = RGBColor(0xC0, 0x00, 0x00)   # red (FAIL)
COLOR_SUBTLE    = RGBColor(0x40, 0x40, 0x40)   # dark grey for body

FONT_BODY       = "Arial"
FONT_CODE       = "Courier New"

# ---------------------------------------------------------------------------
# Document styling helpers
# ---------------------------------------------------------------------------

def _new_doc(title: str, project_name: str = "") -> Document:
    """Create a new Document with consistent base styles."""
    doc = Document()

    # Page margins — A4 with moderate margins
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_SUBTLE

    return doc


def _add_title_block(doc: Document, title: str, subtitle: str = "",
                     project: str = "") -> None:
    """Add a formatted title block at the top of the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_TITLE
    run.font.name = FONT_BODY

    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.font.color.rgb = COLOR_HEADING
        r2.font.name = FONT_BODY

    meta_parts = []
    if project:
        meta_parts.append(f"Project: {project}")
    meta_parts.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_parts.append("Tool: embedded-engineering / doc_output_mcp")

    p3 = doc.add_paragraph(" | ".join(meta_parts))
    p3.runs[0].font.size = Pt(8)
    p3.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Horizontal rule via direct XML (python-docx ≥ 1.x removed get_or_add_pBdr)
    sep = doc.add_paragraph()
    from lxml import etree
    pPr = sep._p.get_or_add_pPr()
    pBdr_xml = (
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="2E74B5"/>'
        '</w:pBdr>'
    )
    pPr.append(etree.fromstring(pBdr_xml))

    doc.add_paragraph()   # spacer


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_TITLE
    run.font.name = FONT_BODY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_HEADING
    run.font.name = FONT_BODY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = FONT_BODY


def _add_kv(doc: Document, key: str, value: str,
            color: Optional[RGBColor] = None) -> None:
    """Add a single key: value line."""
    p = doc.add_paragraph()
    k_run = p.add_run(f"{key}: ")
    k_run.bold = True
    k_run.font.size = Pt(10)
    k_run.font.name = FONT_BODY
    v_run = p.add_run(str(value))
    v_run.font.size = Pt(10)
    v_run.font.name = FONT_BODY
    if color:
        v_run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(1)


def _add_table(doc: Document, headers: List[str],
               rows: List[List[str]]) -> None:
    """Add a formatted table with a header row."""
    if not rows:
        return

    n_cols = len(headers)
    table  = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background (direct XML — python-docx >= 1.x removed get_or_add_shd)
        from lxml import etree as _et
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd_xml = (
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:val="clear" w:color="auto" w:fill="2E74B5"/>'
        )
        for _old_shd in tcPr.findall(qn("w:shd")):
            tcPr.remove(_old_shd)
        tcPr.append(_et.fromstring(shd_xml))

    # Data rows
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            run = cells[i].paragraphs[0].runs[0]
            run.font.name = FONT_BODY
            run.font.size = Pt(9)

    doc.add_paragraph()   # spacer after table


def _add_status_row(doc: Document, rule_id: str, status: str,
                    desc: str, detail: str) -> None:
    """Add a colour-coded rule result line."""
    color_map = {"PASS": COLOR_ACCENT, "WARN": COLOR_WARN, "FAIL": COLOR_FAIL}
    color = color_map.get(status.upper(), COLOR_SUBTLE)
    p = doc.add_paragraph()
    badge = p.add_run(f"[{status}] ")
    badge.bold = True
    badge.font.color.rgb = color
    badge.font.size = Pt(10)
    badge.font.name = FONT_BODY
    id_run = p.add_run(f"{rule_id}  ")
    id_run.font.color.rgb = COLOR_SUBTLE
    id_run.font.size = Pt(10)
    id_run.font.name = FONT_BODY
    desc_run = p.add_run(f"{desc} — ")
    desc_run.bold = True
    desc_run.font.size = Pt(10)
    desc_run.font.name = FONT_BODY
    det_run = p.add_run(detail)
    det_run.font.size = Pt(9)
    det_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    det_run.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(1)


def _ensure_output_dir(output_dir: str) -> Path:
    path = Path(output_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _save(doc: Document, output_dir: str, filename: str) -> str:
    dir_path = _ensure_output_dir(output_dir)
    file_path = dir_path / filename
    doc.save(str(file_path))
    return str(file_path.resolve())


# ---------------------------------------------------------------------------
# Pydantic input models
# ---------------------------------------------------------------------------

class RFReportInput(BaseModel):
    model_config = ConfigDict(
        str_strip_whitespace=True,
        validate_assignment=True,
        extra="forbid",
    )

    project_name: str = Field(
        default="RF Design",
        description="Project name shown in the document header.",
    )
    microstrip_json: Optional[str] = Field(
        default=None,
        description=(
            "JSON string (or file path) containing calc_microstrip.py output. "
            "Paste the full JSON output directly."
        ),
    )
    cpwg_json: Optional[str] = Field(
        default=None,
        description="JSON string (or file path) containing calc_cpwg.py output.",
    )
    antenna_json: Optional[str] = Field(
        default=None,
        description="JSON string (or file path) containing calc_antenna.py output.",
    )
    matching_json: Optional[str] = Field(
        default=None,
        description="JSON string (or file path) containing calc_matching.py output.",
    )
    rules_json: Optional[str] = Field(
        default=None,
        description="JSON string (or file path) containing check_rf_rules.py output.",
    )
    design_notes: str = Field(
        default="",
        description="Optional free-text design notes to include in the document.",
    )
    output_dir: str = Field(
        default=DEFAULT_OUTPUT_DIR,
        description="Directory where the .docx file will be saved.",
    )


class PowerReportInput(BaseModel):
    model_config = ConfigDict(
        str_strip_whitespace=True,
        validate_assignment=True,
        extra="forbid",
    )

    project_name: str = Field(
        default="Embedded System",
        description="Project name shown in the document header.",
    )
    power_tree_json: str = Field(
        ...,
        description="JSON string containing gen_power_tree.py output.",
    )
    mcu_report_json: Optional[str] = Field(
        default=None,
        description="Optional JSON string from gen_mcu_selection_report.py output.",
    )
    design_notes: str = Field(default="")
    output_dir: str = Field(default=DEFAULT_OUTPUT_DIR)


class ProjectSummaryInput(BaseModel):
    model_config = ConfigDict(
        str_strip_whitespace=True,
        validate_assignment=True,
        extra="forbid",
    )

    project_name: str = Field(
        ...,
        description="Project name.",
        min_length=1,
    )
    description: str = Field(
        default="",
        description="Brief project description (1–3 sentences).",
    )
    platform: str = Field(
        default="",
        description="MCU/platform, e.g. 'ESP32-S3', 'STM32G4'.",
    )
    sections: Optional[str] = Field(
        default=None,
        description=(
            "JSON string: list of {title, content} dicts for custom sections. "
            "Example: [{\"title\": \"Architecture\", \"content\": \"...\"}]"
        ),
    )
    rf_design_summary: Optional[str] = Field(
        default=None,
        description="Plain-text or markdown RF design summary to include.",
    )
    power_tree_json: Optional[str] = Field(
        default=None,
        description="JSON from gen_power_tree.py to embed a power section.",
    )
    output_dir: str = Field(default=DEFAULT_OUTPUT_DIR)


class ExportMarkdownInput(BaseModel):
    model_config = ConfigDict(
        str_strip_whitespace=True,
        validate_assignment=True,
        extra="forbid",
    )

    markdown_text: str = Field(
        ...,
        description="Markdown content to convert to Word format.",
        min_length=1,
    )
    document_title: str = Field(
        default="Document",
        description="Title shown at the top of the Word document.",
    )
    output_dir: str = Field(default=DEFAULT_OUTPUT_DIR)
    filename: str = Field(
        default="",
        description=(
            "Output filename (without .docx extension). "
            "If empty, a name is derived from document_title."
        ),
    )


# ---------------------------------------------------------------------------
# JSON loader helper
# ---------------------------------------------------------------------------

def _load_json(raw: Optional[str]) -> Optional[Dict]:
    """Parse a JSON string, returning None if input is absent or invalid."""
    if not raw or not raw.strip():
        return None
    try:
        # Try treating it as a file path first
        p = Path(raw.strip())
        if p.exists() and p.suffix == ".json":
            return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        pass
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return None


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def _build_rf_report(
    project_name: str,
    ms: Optional[Dict], cpwg: Optional[Dict],
    ant: Optional[Dict], match: Optional[Dict],
    rules: Optional[Dict], notes: str,
) -> Document:
    doc = _new_doc("RF Design Report", project_name)
    _add_title_block(doc, "RF Design Report", project_name)

    # ---- Microstrip ----
    if ms and ms.get("success"):
        _add_h1(doc, "1. Microstrip Feedline")
        r = ms.get("results", {})
        inp = ms.get("inputs", {})
        rows = [
            ["Substrate", inp.get("substrate", "—")],
            ["εr", str(inp.get("er", "—"))],
            ["Substrate height h", f"{inp.get('h_mm', '—')} mm"],
            ["Target Z₀", f"{inp.get('target_z0', 50)} Ω"],
            ["Calculated width", f"{r.get('width_mm', '—')} mm"],
            ["Actual Z₀", f"{r.get('z0_ohm', '—')} Ω"],
            ["Effective εr", str(r.get('effective_er', '—'))],
            ["Guided wavelength λg", f"{r.get('guided_wavelength_mm', '—')} mm"],
            ["λ/4 length", f"{r.get('quarter_wave_mm', '—')} mm"],
            ["Total insertion loss", f"{r.get('total_loss_db_cm', '—')} dB/cm"],
        ]
        _add_table(doc, ["Parameter", "Value"], rows)
        for w in ms.get("warnings", []):
            _add_body(doc, f"⚠ {w}")

    # ---- CPWG ----
    if cpwg and cpwg.get("success"):
        _add_h1(doc, "2. CPWG Feedline")
        r = cpwg.get("results", {})
        inp = cpwg.get("inputs", {})
        rows = [
            ["Signal width w", f"{inp.get('width_mm', '—')} mm"],
            ["Gap g", f"{inp.get('gap_mm', '—')} mm"],
            ["Z₀", f"{r.get('z0_ohm', '—')} Ω"],
            ["Effective εr", str(r.get('effective_er', '—'))],
            ["λg", f"{r.get('guided_wavelength_mm', '—')} mm"],
            ["λ/4 length", f"{r.get('quarter_wave_mm', '—')} mm"],
            ["Max via fence pitch", f"{r.get('via_fence_pitch_mm', '—')} mm"],
            ["Min board-edge clearance", f"{r.get('recommended_edge_clearance_mm', '—')} mm"],
        ]
        _add_table(doc, ["Parameter", "Value"], rows)
        for w in cpwg.get("warnings", []):
            _add_body(doc, f"⚠ {w}")

    # ---- Antenna ----
    if ant and ant.get("success"):
        _add_h1(doc, "3. Antenna Design")
        atype = ant.get("inputs", {}).get("antenna_type", "—").title()
        freq  = ant.get("inputs", {}).get("freq_ghz", "—")
        _add_body(doc, f"Type: {atype}  |  Frequency: {freq} GHz")
        r = ant.get("results", {})
        rows = [[k.replace("_", " ").title(), str(v)]
                for k, v in r.items()
                if not isinstance(v, (dict, list)) and v is not None]
        _add_table(doc, ["Parameter", "Value"], rows)
        for rec in ant.get("recommendations", [])[:4]:
            _add_body(doc, f"• {rec}")

    # ---- Matching ----
    if match and match.get("success"):
        _add_h1(doc, "4. Matching Network")
        ntype = match.get("inputs", {}).get("network_type", "—").upper()
        _add_body(doc, f"Topology: {ntype}")
        r = match.get("results", {})

        def _add_topo(topo_dict: Optional[Dict], label: str) -> None:
            if not topo_dict:
                return
            _add_h2(doc, label)
            for k, v in topo_dict.items():
                if isinstance(v, dict) and "type" in v:
                    ctype = v.get("type", "")
                    val   = (f"L = {v.get('L_nH')} nH"
                             if ctype == "inductor"
                             else f"C = {v.get('C_pF')} pF")
                    _add_body(doc, f"  {k}: {ctype.title()} — {val}")
                elif not isinstance(v, (dict, list)):
                    _add_body(doc, f"  {k}: {v}")

        if ntype == "STUB":
            for sol in r.get("solutions", []):
                _add_h2(doc, f"Solution {sol.get('solution')}")
                d = sol.get("distance_from_load", {})
                _add_body(doc, f"  Distance from load: {d.get('physical_length_mm')} mm "
                               f"({d.get('electrical_length_deg')}°)")
                for stub_type, stub_data in sol.get("stub_options", {}).items():
                    _add_body(doc, f"  {stub_type.replace('_', ' ').title()}: "
                                   f"{stub_data.get('physical_length_mm')} mm "
                                   f"({stub_data.get('electrical_length_deg')}°)")
        else:
            Q = r.get("Q") or r.get("Q_target")
            if Q:
                _add_body(doc, f"  Network Q = {Q}")
            _add_topo(r.get("low_pass"), "Low-pass topology")
            _add_topo(r.get("high_pass"), "High-pass topology")

    # ---- RF Rules ----
    if rules and rules.get("success"):
        _add_h1(doc, "5. RF PCB Rule Check")
        s = rules.get("summary", {})
        overall = s.get("overall_status", "—")
        color   = {"PASS": COLOR_ACCENT, "WARN": COLOR_WARN, "FAIL": COLOR_FAIL}.get(
            overall, COLOR_SUBTLE)
        _add_kv(doc, "Overall status", overall, color)
        _add_kv(doc, "Rules evaluated",
                f"{s.get('total_rules','—')} total  —  "
                f"PASS: {s.get('pass_count','—')}  "
                f"WARN: {s.get('warn_count','—')}  "
                f"FAIL: {s.get('fail_count','—')}")
        doc.add_paragraph()

        for rule in rules.get("rules", []):
            _add_status_row(
                doc,
                rule.get("rule_id", ""),
                rule.get("status", ""),
                rule.get("description", ""),
                rule.get("detail", ""),
            )
            if rule.get("status") in ("WARN", "FAIL"):
                rec_p = doc.add_paragraph(f"     → {rule.get('recommendation','')}")
                rec_p.runs[0].font.size = Pt(9)
                rec_p.runs[0].font.color.rgb = RGBColor(0x50, 0x50, 0x50)
                rec_p.runs[0].font.name = FONT_BODY

    # ---- Notes ----
    if notes.strip():
        _add_h1(doc, "Design Notes")
        _add_body(doc, notes)

    return doc


def _build_power_report(
    project_name: str,
    pt: Dict,
    mcu: Optional[Dict],
    notes: str,
) -> Document:
    doc = _new_doc("Power Architecture Report", project_name)
    _add_title_block(doc, "Power Architecture Report", project_name)

    r = pt.get("results", {})

    _add_h1(doc, "1. Power Tree")
    for line in r.get("power_tree_text", []):
        p = doc.add_paragraph(line)
        p.runs[0].font.name = FONT_CODE
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()

    _add_h1(doc, "2. System Power Summary")
    rows = [
        ["Total load power",   f"{r.get('total_load_mw', '—')} mW"],
        ["Regulator losses",   f"{r.get('total_dissipation_mw', '—')} mW"],
        ["Total input power",  f"{r.get('total_input_power_mw', '—')} mW"],
        ["System efficiency",  f"{r.get('system_efficiency_pct', '—')} %"],
    ]
    _add_table(doc, ["Metric", "Value"], rows)

    _add_h1(doc, "3. Per-Rail Analysis")
    for rail in r.get("rails", []):
        _add_h2(doc, f"{rail.get('name')} — {rail.get('voltage_v')} V")
        rail_rows = [
            ["Regulator type",   rail.get("regulator_type", "—").upper()],
            ["Load current",     f"{rail.get('current_ma', '—')} mA"],
            ["Load power",       f"{rail.get('power_mw', '—')} mW"],
            ["Efficiency est.",  f"{rail.get('efficiency_pct', '—')} %"],
            ["Dissipation est.", f"{rail.get('dissipation_mw', '—')} mW"],
            ["Current headroom", f"{rail.get('headroom_pct', '—')} %"],
            ["LDO dropout OK",   str(rail.get('dropout_ok', '—'))],
            ["Consumers",        ", ".join(rail.get("consumers", []))],
        ]
        _add_table(doc, ["Parameter", "Value"], rail_rows)
        for issue in rail.get("issues", []):
            _add_body(doc, f"⚠ {issue}")
        for bom in rail.get("bom_hints", [])[:2]:
            _add_body(doc, f"• {bom}")

    if r.get("sequencing_notes"):
        _add_h1(doc, "4. Sequencing Notes")
        for note in r["sequencing_notes"]:
            _add_body(doc, f"• {note}")

    if mcu and mcu.get("success"):
        _add_h1(doc, "5. MCU Selection Summary")
        mr = mcu.get("results", {})
        top = mr.get("top_candidates", [])
        if top:
            rows = [
                [str(c.get("rank")), c.get("part", ""), c.get("family", ""),
                 str(c.get("score", ""))]
                for c in top
            ]
            _add_table(doc, ["Rank", "Part", "Family", "Score"], rows)
            first = top[0]
            _add_body(doc, f"Primary recommendation: {first.get('part')} — "
                           f"{first.get('rationale', '')[:200]}")

    if notes.strip():
        _add_h1(doc, "Design Notes")
        _add_body(doc, notes)

    return doc


def _markdown_to_doc(doc: Document, md: str) -> None:
    """Convert markdown text into Word paragraphs (best-effort)."""
    for line in md.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            _add_h2(doc, stripped[4:])
        elif stripped.startswith("## "):
            _add_h1(doc, stripped[3:])
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = COLOR_TITLE
            r.font.name = FONT_BODY
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(stripped[2:]).font.name = FONT_BODY
        elif stripped.startswith("|"):
            # Rough table row — just add as code-style paragraph
            p = doc.add_paragraph(stripped)
            p.runs[0].font.name = FONT_CODE
            p.runs[0].font.size = Pt(9)
        elif stripped.startswith("```"):
            pass   # skip fence markers
        elif stripped == "---" or stripped == "===":
            pass   # skip horizontal rules
        elif stripped == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(stripped)
            if p.runs:
                p.runs[0].font.name = FONT_BODY
                p.runs[0].font.size = Pt(10)


def _safe_filename(name: str) -> str:
    """Convert a string to a safe filename component."""
    name = re.sub(r"[^\w\s-]", "", name).strip()
    name = re.sub(r"\s+", "_", name)
    return name[:60]


# ---------------------------------------------------------------------------
# Tools
# ---------------------------------------------------------------------------

@mcp.tool(
    name="doc_rf_design_report",
    annotations={
        "title": "Generate RF Design Report (Word)",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": False,
        "openWorldHint": False,
    },
)
async def doc_rf_design_report(params: RFReportInput) -> str:
    """Generate a professional RF design report (.docx) from calc_* script outputs.

    Assembles microstrip, CPWG, antenna, matching network, and RF rule-check
    results into a formatted Word document with colour-coded pass/warn/fail
    indicators. Pass the JSON output from each script directly as a string.

    Args:
        params (RFReportInput): Report parameters including:
            - project_name (str): Project name for the document header
            - microstrip_json (Optional[str]): JSON from calc_microstrip.py
            - cpwg_json (Optional[str]): JSON from calc_cpwg.py
            - antenna_json (Optional[str]): JSON from calc_antenna.py
            - matching_json (Optional[str]): JSON from calc_matching.py
            - rules_json (Optional[str]): JSON from check_rf_rules.py
            - design_notes (str): Optional free-text notes
            - output_dir (str): Directory to save the .docx file

    Returns:
        str: Absolute path to the generated .docx file, or an error message.
    """
    try:
        ms    = _load_json(params.microstrip_json)
        cpwg  = _load_json(params.cpwg_json)
        ant   = _load_json(params.antenna_json)
        match = _load_json(params.matching_json)
        rules = _load_json(params.rules_json)

        if not any([ms, cpwg, ant, match, rules]):
            return json.dumps({
                "error": "No valid JSON data provided. "
                         "Pass at least one calc_* script output as a JSON string."
            })

        doc = _build_rf_report(
            params.project_name, ms, cpwg, ant, match, rules,
            params.design_notes,
        )

        ts  = datetime.now().strftime("%Y%m%d_%H%M")
        fname = f"RF_Report_{_safe_filename(params.project_name)}_{ts}.docx"
        path  = _save(doc, params.output_dir, fname)
        return json.dumps({"success": True, "file_path": path, "filename": fname})

    except Exception as exc:
        return json.dumps({"error": f"{type(exc).__name__}: {exc}"})


@mcp.tool(
    name="doc_power_tree_report",
    annotations={
        "title": "Generate Power Architecture Report (Word)",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": False,
        "openWorldHint": False,
    },
)
async def doc_power_tree_report(params: PowerReportInput) -> str:
    """Generate a power architecture report (.docx) from gen_power_tree output.

    Produces a document with the power tree diagram, per-rail efficiency and
    thermal analysis, BOM hints, and optionally an MCU selection summary.

    Args:
        params (PowerReportInput): Report parameters including:
            - project_name (str): Project name
            - power_tree_json (str): JSON output from gen_power_tree.py
            - mcu_report_json (Optional[str]): JSON from gen_mcu_selection_report.py
            - design_notes (str): Optional notes
            - output_dir (str): Output directory

    Returns:
        str: Absolute path to the generated .docx file, or an error message.
    """
    try:
        pt  = _load_json(params.power_tree_json)
        mcu = _load_json(params.mcu_report_json)

        if not pt or not pt.get("success"):
            return json.dumps({
                "error": "power_tree_json is required and must be valid output from gen_power_tree.py."
            })

        doc  = _build_power_report(params.project_name, pt, mcu, params.design_notes)
        ts   = datetime.now().strftime("%Y%m%d_%H%M")
        fname = f"Power_Report_{_safe_filename(params.project_name)}_{ts}.docx"
        path  = _save(doc, params.output_dir, fname)
        return json.dumps({"success": True, "file_path": path, "filename": fname})

    except Exception as exc:
        return json.dumps({"error": f"{type(exc).__name__}: {exc}"})


@mcp.tool(
    name="doc_project_summary",
    annotations={
        "title": "Generate Full Project Summary Document (Word)",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": False,
        "openWorldHint": False,
    },
)
async def doc_project_summary(params: ProjectSummaryInput) -> str:
    """Generate a complete project summary document (.docx).

    Creates a structured project overview document with description, platform
    info, custom sections, RF design summary, and power tree data.
    Useful as a project handover document or design review report.

    Args:
        params (ProjectSummaryInput): Document parameters including:
            - project_name (str): Project name (required)
            - description (str): Brief project description
            - platform (str): MCU/SoC platform
            - sections (Optional[str]): JSON list of {title, content} dicts
            - rf_design_summary (Optional[str]): RF summary text or markdown
            - power_tree_json (Optional[str]): JSON from gen_power_tree.py
            - output_dir (str): Output directory

    Returns:
        str: Absolute path to the generated .docx file, or an error message.
    """
    try:
        doc = _new_doc("Project Summary", params.project_name)
        _add_title_block(
            doc, params.project_name,
            "Project Design Summary",
            params.project_name,
        )

        # Overview
        _add_h1(doc, "1. Project Overview")
        if params.description:
            _add_body(doc, params.description)
        if params.platform:
            _add_kv(doc, "Platform", params.platform)
        _add_kv(doc, "Document date", datetime.now().strftime("%Y-%m-%d"))

        # Custom sections
        section_num = 2
        custom = _load_json(params.sections)
        if isinstance(custom, list):
            for sec in custom:
                if isinstance(sec, dict):
                    _add_h1(doc, f"{section_num}. {sec.get('title', 'Section')}")
                    _markdown_to_doc(doc, str(sec.get("content", "")))
                    section_num += 1

        # RF summary
        if params.rf_design_summary:
            _add_h1(doc, f"{section_num}. RF Design")
            _markdown_to_doc(doc, params.rf_design_summary)
            section_num += 1

        # Power tree
        pt = _load_json(params.power_tree_json)
        if pt and pt.get("success"):
            _add_h1(doc, f"{section_num}. Power Architecture")
            r = pt.get("results", {})
            for line in r.get("power_tree_text", []):
                p = doc.add_paragraph(line)
                if p.runs:
                    p.runs[0].font.name = FONT_CODE
                    p.runs[0].font.size = Pt(9)
                p.paragraph_format.space_after = Pt(1)
            doc.add_paragraph()
            _add_kv(doc, "System efficiency",
                    f"{r.get('system_efficiency_pct', '—')} %")
            _add_kv(doc, "Total load power",
                    f"{r.get('total_load_mw', '—')} mW")
            section_num += 1

        ts   = datetime.now().strftime("%Y%m%d_%H%M")
        fname = f"Summary_{_safe_filename(params.project_name)}_{ts}.docx"
        path  = _save(doc, params.output_dir, fname)
        return json.dumps({"success": True, "file_path": path, "filename": fname})

    except Exception as exc:
        return json.dumps({"error": f"{type(exc).__name__}: {exc}"})


@mcp.tool(
    name="doc_export_markdown",
    annotations={
        "title": "Export Markdown as Word Document",
        "readOnlyHint": False,
        "destructiveHint": False,
        "idempotentHint": False,
        "openWorldHint": False,
    },
)
async def doc_export_markdown(params: ExportMarkdownInput) -> str:
    """Convert any markdown text into a formatted Word document (.docx).

    Useful for exporting workflow documents, design notes, or any text
    content that Claude has generated in markdown format.

    Args:
        params (ExportMarkdownInput): Export parameters including:
            - markdown_text (str): Markdown content to convert
            - document_title (str): Title for the document header
            - output_dir (str): Output directory
            - filename (str): Output filename (without .docx); auto-derived if empty

    Returns:
        str: Absolute path to the generated .docx file, or an error message.
    """
    try:
        doc = _new_doc(params.document_title)
        _add_title_block(doc, params.document_title)
        _markdown_to_doc(doc, params.markdown_text)

        base = (params.filename.strip()
                or _safe_filename(params.document_title))
        ts   = datetime.now().strftime("%Y%m%d_%H%M")
        fname = f"{base}_{ts}.docx"
        path  = _save(doc, params.output_dir, fname)
        return json.dumps({"success": True, "file_path": path, "filename": fname})

    except Exception as exc:
        return json.dumps({"error": f"{type(exc).__name__}: {exc}"})


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    mcp.run()
